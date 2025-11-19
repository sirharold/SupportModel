#!/usr/bin/env python3
"""
Agrega números de página VISIBLES en el pie de página
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

base_dir = Path("/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Octubre2025")
words_dir = base_dir / "Words"
doc_path = words_dir / "ProyectoMagisterHaroldGomez.docx"

def add_page_number_to_footer(section):
    """Agrega número de página visible en el pie de página"""
    # Obtener o crear footer
    footer = section.footer

    # Limpiar contenido existente del footer
    for element in footer._element:
        if element.tag.endswith('p'):
            footer._element.remove(element)

    # Crear nuevo párrafo centrado para el número de página
    footer_para = footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar run para el número de página
    run = footer_para.add_run()

    # Crear campo de número de página
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

print("AGREGANDO NÚMEROS DE PÁGINA VISIBLES")
print("=" * 70)

# Cargar documento
print(f"\n1. Cargando: {doc_path.name}")
doc = Document(str(doc_path))

print(f"   Secciones encontradas: {len(doc.sections)}")

# Agregar números de página a cada sección
print(f"\n2. Agregando números de página visibles...")

for i, section in enumerate(doc.sections):
    print(f"   - Sección {i}: Agregando número de página en pie de página")
    add_page_number_to_footer(section)

# Guardar
print(f"\n3. Guardando documento...")
doc.save(str(doc_path))

print(f"\n{'='*70}")
print(f"✅ NÚMEROS DE PÁGINA AGREGADOS")
print(f"{'='*70}")
print(f"   Se agregaron números de página visibles a {len(doc.sections)} secciones")
print(f"   Los números se muestran centrados en el pie de página")
print(f"{'='*70}")
