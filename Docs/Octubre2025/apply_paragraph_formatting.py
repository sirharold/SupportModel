#!/usr/bin/env python3
"""
Aplica sangría de primera línea (1.27 cm) a párrafos normales
IMPORTANTE: NO aplica sangría a párrafos dentro de tablas
"""

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

base_dir = Path("/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Octubre2025")
words_dir = base_dir / "Words"
doc_path = words_dir / "ProyectoMagisterHaroldGomez.docx"

# Constante de sangría
SANGRIA_PRIMERA_LINEA = Cm(1.27)

# Estilos que NO deben tener sangría
EXCLUDED_STYLES = [
    'Title',
    'Subtitle',
    'Heading 1',
    'Heading 2',
    'Heading 3',
    'Heading 4',
    'Heading 5',
    'Heading 6',
    'toc',
    'TOC',
    'Caption',
    'Table',
    'Image Caption',
    'Figure'
]

def get_all_table_paragraphs(doc):
    """Obtiene todos los párrafos que están dentro de tablas"""
    table_paragraphs = set()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Usar el objeto _element como identificador único
                    table_paragraphs.add(id(paragraph._element))

    return table_paragraphs

print("APLICANDO FORMATO DE PÁRRAFOS")
print("=" * 70)

# Cargar documento
print(f"\n1. Cargando: {doc_path.name}")
doc = Document(str(doc_path))

# Obtener párrafos que están en tablas
print(f"2. Identificando párrafos en tablas...")
table_para_ids = get_all_table_paragraphs(doc)
print(f"   Párrafos en tablas: {len(table_para_ids)}")

# Aplicar formato
print(f"3. Aplicando sangría de {SANGRIA_PRIMERA_LINEA.cm} cm a párrafos normales...")

stats = {
    'total': len(doc.paragraphs),
    'en_tabla': 0,
    'estilo_excluido': 0,
    'vacio': 0,
    'con_sangria': 0
}

for para in doc.paragraphs:
    # Contar total
    para_id = id(para._element)

    # Saltar párrafos vacíos
    if not para.text.strip():
        stats['vacio'] += 1
        continue

    # Saltar párrafos en tablas
    if para_id in table_para_ids:
        stats['en_tabla'] += 1
        continue

    # Saltar estilos excluidos (sin acceder a para.style para evitar error de relaciones duplicadas)
    try:
        style_name = para.style.name if para.style else ''
        if any(excluded in style_name for excluded in EXCLUDED_STYLES):
            stats['estilo_excluido'] += 1
            continue
    except (ValueError, KeyError):
        # Si no podemos acceder al estilo, verificar por texto
        # Si el párrafo parece un título (corto, sin punto final), excluirlo
        if len(para.text.strip()) < 100 and not para.text.strip().endswith('.'):
            stats['estilo_excluido'] += 1
            continue

    # Aplicar sangría de primera línea
    para.paragraph_format.first_line_indent = SANGRIA_PRIMERA_LINEA
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    stats['con_sangria'] += 1

# Guardar
print(f"4. Guardando documento...")
doc.save(str(doc_path))

# Resumen
print(f"\n{'='*70}")
print(f"✅ FORMATO APLICADO")
print(f"{'='*70}")
print(f"   Total de párrafos: {stats['total']}")
print(f"   └─ Vacíos: {stats['vacio']}")
print(f"   └─ En tablas (sin sangría): {stats['en_tabla']}")
print(f"   └─ Estilos excluidos: {stats['estilo_excluido']}")
print(f"   └─ Con sangría aplicada: {stats['con_sangria']}")
print(f"\n   Sangría: {SANGRIA_PRIMERA_LINEA.cm} cm en primera línea")
print(f"   Alineación: Justificado")
print(f"{'='*70}")
