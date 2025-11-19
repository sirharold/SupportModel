#!/usr/bin/env python3
"""
Convierte todo el documento en un solo paso con pandoc
Esto preserva TODAS las imágenes
"""

from pathlib import Path
import subprocess

base_dir = Path("/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Octubre2025")
temp_dir = base_dir / "temp"
words_dir = base_dir / "Words"

print("CREANDO DOCUMENTO COMPLETO EN UN SOLO PASO")
print("=" * 70)

# 1. Combinar todo el contenido markdown en un solo archivo
print(f"\n1. Creando archivo markdown completo...")

combined_md = temp_dir / "documento_completo.md"

with open(combined_md, 'w', encoding='utf-8') as outfile:
    # Preliminares (Resumen + Abstract) con salto de sección
    prelim_md = temp_dir / "preliminares_combined.md"
    with open(prelim_md, 'r', encoding='utf-8') as f:
        outfile.write(f.read())

    # Salto de sección antes de capítulos (para cambiar numeración)
    outfile.write('\n\n```{=openxml}\n<w:p><w:pPr><w:sectPr>\n')
    outfile.write('  <w:pgNumType w:fmt="decimal" w:start="1"/>\n')
    outfile.write('</w:sectPr></w:pPr></w:p>\n```\n\n')

    # Capítulos (1-8 + Bibliografía)
    cap_md = temp_dir / "capitulos_combined.md"
    with open(cap_md, 'r', encoding='utf-8') as f:
        outfile.write(f.read())

print(f"   ✓ Creado: {combined_md.name}")

# 2. Convertir a Word en un solo paso
print(f"\n2. Convirtiendo a Word con pandoc...")

output_docx = temp_dir / "documento_completo.docx"

cmd = [
    'pandoc',
    str(combined_md),
    '-o', str(output_docx),
    f'--reference-doc={base_dir}/Plantilla_Tesis_2025.dotx',
    '--resource-path=.'
]

subprocess.run(cmd, cwd=base_dir, check=True)

print(f"   ✓ Creado: {output_docx.name}")

# 3. Ahora combinar First.docx con el documento completo
print(f"\n3. Combinando First.docx con documento completo...")

from docx import Document
from docxcompose.composer import Composer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_page_number_format(section, format_type='decimal', start=1):
    """Establece el formato de numeración de páginas para una sección"""
    sectPr = section._sectPr
    for child in list(sectPr):
        if child.tag == qn('w:pgNumType'):
            sectPr.remove(child)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), format_type)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)

# Cargar First.docx
first_path = words_dir / "First.docx"
master = Document(str(first_path))

# Configurar numeración romana para First
for section in master.sections:
    set_page_number_format(section, format_type='lowerRoman', start=1)

# Agregar documento completo
composer = Composer(master)
doc_completo = Document(str(output_docx))
composer.append(doc_completo)

# Guardar resultado final
final_path = words_dir / "ProyectoMagisterHaroldGomez.docx"
print(f"\n4. Guardando: {final_path.name}")
composer.save(str(final_path))

# Verificación
from docx import Document
result = Document(str(final_path))
image_count = sum(1 for rel in result.part.rels.values() if "image" in rel.target_ref)

print(f"\n{'='*70}")
print(f"✅ DOCUMENTO FINAL CREADO")
print(f"{'='*70}")
print(f"   Secciones: {len(result.sections)}")
print(f"   Párrafos: {len(result.paragraphs)}")
print(f"   Tablas: {len(result.tables)}")
print(f"   Imágenes embebidas: {image_count}")
print(f"   Tamaño: {final_path.stat().st_size / (1024*1024):.1f} MB")
print(f"\n   Numeración de páginas:")
print(f"   - First + preliminares: Números romanos (i, ii, iii...)")
print(f"   - Capítulos 1-8 + Bibliografía: Números arábigos (1, 2, 3...)")
print(f"{'='*70}")
