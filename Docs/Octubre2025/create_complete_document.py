#!/usr/bin/env python3
"""
Crea el documento completo en un solo paso con pandoc
Combina: First.md + preliminares + capítulos
"""

from pathlib import Path
import subprocess

base_dir = Path("/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Octubre2025")
temp_dir = base_dir / "temp"
words_dir = base_dir / "Words"

print("CREANDO DOCUMENTO COMPLETO EN UN SOLO PASO CON PANDOC")
print("=" * 70)

# 1. Combinar First + preliminares + capítulos en un solo markdown
print(f"\n1. Combinando todo el contenido markdown...")

full_md = temp_dir / "full_document.md"

with open(full_md, 'w', encoding='utf-8') as outfile:
    # First (portada)
    first_md = temp_dir / "First.md"
    with open(first_md, 'r', encoding='utf-8') as f:
        content = f.read()
        # Limpiar contenido de First - solo queremos el texto visual
        outfile.write(content)

    # Salto de página después de First
    outfile.write('\n\n```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```\n\n')

    # Preliminares (Resumen + Abstract) - con numeración romana
    prelim_md = temp_dir / "preliminares_combined.md"
    with open(prelim_md, 'r', encoding='utf-8') as f:
        outfile.write(f.read())

    # Salto de sección para cambiar a numeración arábiga
    outfile.write('\n\n```{=openxml}\n')
    outfile.write('<w:p>\n')
    outfile.write('  <w:pPr>\n')
    outfile.write('    <w:sectPr>\n')
    outfile.write('      <w:pgNumType w:fmt="decimal" w:start="1"/>\n')
    outfile.write('    </w:sectPr>\n')
    outfile.write('  </w:pPr>\n')
    outfile.write('</w:p>\n')
    outfile.write('```\n\n')

    # Capítulos (1-8 + Bibliografía) - con numeración arábiga
    cap_md = temp_dir / "capitulos_combined.md"
    with open(cap_md, 'r', encoding='utf-8') as f:
        outfile.write(f.read())

print(f"   ✓ Creado: {full_md.name}")

# 2. Convertir todo a Word en un solo paso
print(f"\n2. Convirtiendo a Word con pandoc (esto preserva TODAS las imágenes)...")

output_docx = words_dir / "ProyectoMagisterHaroldGomez.docx"

cmd = [
    'pandoc',
    str(full_md),
    '-o', str(output_docx),
    f'--reference-doc={base_dir}/Plantilla_Tesis_2025.dotx',
    '--resource-path=.'
]

subprocess.run(cmd, cwd=base_dir, check=True)

print(f"   ✓ Creado: {output_docx.name}")

# Verificación
from docx import Document

doc = Document(str(output_docx))
image_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.target_ref)

print(f"\n{'='*70}")
print(f"✅ DOCUMENTO FINAL CREADO")
print(f"{'='*70}")
print(f"   Secciones: {len(doc.sections)}")
print(f"   Párrafos: {len(doc.paragraphs)}")
print(f"   Tablas: {len(doc.tables)}")
print(f"   Imágenes embebidas: {image_count}")
print(f"   Tamaño: {output_docx.stat().st_size / (1024*1024):.1f} MB")
print(f"\n   Numeración de páginas:")
print(f"   - First + preliminares: Números romanos (i, ii, iii...)")
print(f"   - Capítulos 1-8 + Bibliografía: Números arábigos (1, 2, 3...)")
print(f"{'='*70}")
