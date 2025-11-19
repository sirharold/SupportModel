#!/usr/bin/env python3
"""
Agrega bordes a todas las tablas del documento Word
"""

from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

def add_table_borders(table):
    """
    Agrega bordes a una tabla de Word
    """
    tbl = table._element
    tblPr = tbl.tblPr

    # Si no existe tblPr, crearlo
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remover bordes existentes si los hay
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)

    # Crear nuevo elemento de bordes
    tblBorders = OxmlElement('w:tblBorders')

    # Definir cada tipo de borde
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Grosor: 4 (0.5 pt)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Negro
        tblBorders.append(border)

    tblPr.append(tblBorders)

base_dir = Path("/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Octubre2025")
words_dir = base_dir / "Words"
doc_path = words_dir / "ProyectoMagisterHaroldGomez.docx"

print("AGREGANDO BORDES A TABLAS")
print("=" * 70)

# Cargar documento
print(f"\n1. Cargando: {doc_path.name}")
doc = Document(str(doc_path))

# Contar tablas
total_tables = len(doc.tables)
print(f"2. Tablas encontradas: {total_tables}")

# Agregar bordes a cada tabla
print(f"3. Agregando bordes a todas las tablas...")
for i, table in enumerate(doc.tables, 1):
    add_table_borders(table)
    if i % 5 == 0:
        print(f"   Procesadas: {i}/{total_tables}")

# Guardar documento
print(f"\n4. Guardando documento...")
doc.save(str(doc_path))

print(f"\n{'='*70}")
print(f"✅ BORDES AGREGADOS")
print(f"{'='*70}")
print(f"   Tablas procesadas: {total_tables}")
print(f"   Documento: {doc_path.name}")
print(f"{'='*70}")
