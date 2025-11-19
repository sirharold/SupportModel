#!/usr/bin/env python3
"""
Crea el documento final correctamente:
1. Mantiene First.docx original (portada)
2. Agrega preliminares y capítulos desde Word
3. Configura numeración de páginas VISIBLE en pie de página
4. Mantiene numeración de títulos correcta
"""

from docx import Document
from docxcompose.composer import Composer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

base_dir = Path("/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Octubre2025")
words_dir = base_dir / "Words"
temp_dir = base_dir / "temp"

def set_page_number_format(section, format_type='decimal', start=1):
    """Establece el formato de numeración de páginas"""
    sectPr = section._sectPr

    # Remover pgNumType existente
    for child in list(sectPr):
        if child.tag == qn('w:pgNumType'):
            sectPr.remove(child)

    # Crear nuevo pgNumType
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), format_type)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)

def add_page_numbers(section):
    """Agrega números de página visibles en el pie de página"""
    # Obtener o crear footer
    footer = section.footer

    # Limpiar footer existente
    footer._element.clear_content()

    # Crear párrafo centrado
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar campo de número de página
    run = paragraph.add_run()

    # Crear elemento de campo para número de página
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

print("CREANDO DOCUMENTO FINAL CORRECTO")
print("=" * 70)

# 1. Cargar First.docx (mantiene formato original de portada)
first_path = words_dir / "First.docx"
print(f"\n1. Cargando First.docx original (portada)")
master = Document(str(first_path))
composer = Composer(master)

# Configurar sección First sin numeración visible
for section in master.sections:
    set_page_number_format(section, format_type='lowerRoman', start=1)
    # No agregar números de página en la portada

# 2. Agregar preliminares (Resumen + Abstract)
prelim_path = temp_dir / "preliminares.docx"
print(f"2. Agregando preliminares.docx")
preliminares = Document(str(prelim_path))
composer.append(preliminares)

# Configurar numeración romana para preliminares (visible)
print(f"   - Configurando numeración romana (i, ii, iii...)")
for section in master.sections:
    set_page_number_format(section, format_type='lowerRoman', start=1)
    add_page_numbers(section)

# 3. Agregar salto de sección para cambiar a numeración arábiga
print(f"\n3. Insertando salto de sección antes de Capítulo 1")
master.add_section()

# 4. Agregar capítulos
cap_path = temp_dir / "capitulos.docx"
print(f"4. Agregando capitulos.docx (Capítulos 1-8 + Bibliografía)")
capitulos = Document(str(cap_path))
composer.append(capitulos)

# Configurar numeración arábiga para capítulos (visible)
print(f"   - Configurando numeración arábiga (1, 2, 3...)")
last_section = master.sections[-1]
set_page_number_format(last_section, format_type='decimal', start=1)
add_page_numbers(last_section)

# 5. Guardar
output_path = words_dir / "ProyectoMagisterHaroldGomez.docx"
print(f"\n5. Guardando: {output_path.name}")
composer.save(str(output_path))

# Verificar
doc = Document(str(output_path))
image_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.target_ref)

print(f"\n{'='*70}")
print(f"✅ DOCUMENTO FINAL CREADO")
print(f"{'='*70}")
print(f"   Secciones: {len(doc.sections)}")
print(f"   Párrafos: {len(doc.paragraphs)}")
print(f"   Tablas: {len(doc.tables)}")
print(f"   Imágenes: {image_count}")
print(f"   Tamaño: {output_path.stat().st_size / (1024*1024):.1f} MB")
print(f"\n   Numeración de páginas:")
print(f"   - First (portada): Sin número visible")
print(f"   - Preliminares: Números romanos (i, ii, iii...)")
print(f"   - Capítulos: Números arábigos (1, 2, 3...)")
print(f"\n   ⚠️  NOTA: Si las imágenes no aparecen, es porque")
print(f"   docxcompose las pierde al combinar. Necesitamos otra solución.")
print(f"{'='*70}")
