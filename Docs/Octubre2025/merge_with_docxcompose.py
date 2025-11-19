#!/usr/bin/env python3
"""
Script para combinar documentos Word usando docxcompose
Esta librería maneja correctamente imágenes, tablas y otros elementos complejos
"""

from pathlib import Path
from datetime import datetime
from docx import Document
from docxcompose.composer import Composer
import zipfile

# Configuración
WORDS_DIR = Path(__file__).parent / "Words"
OUTPUT_FILE = WORDS_DIR / "ProyectoMAgisterHaroldGomez.docx"

# Archivos a combinar en orden
FILES_TO_MERGE = [
    'First.docx',
    'Capitulo_0_Resumen.docx',
    'Capitulo_0_Abstract.docx',
    'Capitulo_1_Introduccion.docx',
    'Capitulo_2_Estado_del_Arte.docx',
    'Capitulo_3_Marco_Teorico.docx',
    'Capitulo_4_Analisis_Exploratorio.docx',
    'Capitulo_5_Metodologia.docx',
    'Capitulo_6_Implementacion.docx',
    'Capitulo_7_Resultados.docx',
    'Capitulo_8_Conclusiones_y_Trabajo_Futuro.docx',
    'Bibliografia.docx'
]


def count_images_in_docx(docx_path):
    """Cuenta las imágenes en un archivo Word."""
    try:
        with zipfile.ZipFile(str(docx_path), 'r') as zip_ref:
            files = zip_ref.namelist()
            media_files = [f for f in files if 'word/media/' in f and
                          f.endswith(('.png', '.jpg', '.jpeg', '.gif'))]
            return len(media_files)
    except:
        return 0


def merge_documents(files, output_file):
    """
    Combina múltiples documentos Word usando docxcompose.
    """
    try:
        # Crear compositor con el primer documento
        print(f"\n1. Cargando documento base: {files[0].name}")
        master = Document(files[0])
        composer = Composer(master)

        # Agregar cada documento subsecuente
        for i, file_path in enumerate(files[1:], 2):
            img_count = count_images_in_docx(file_path)
            img_info = f" ({img_count} imágenes)" if img_count > 0 else ""
            print(f"{i}. Agregando: {file_path.name}{img_info}")

            # Cargar y agregar el documento
            doc = Document(file_path)
            composer.append(doc)

        # Guardar
        print(f"\nGuardando documento combinado...")
        composer.save(str(output_file))

        # Verificar
        if output_file.exists():
            size_mb = output_file.stat().st_size / (1024 * 1024)
            total_images = count_images_in_docx(output_file)

            print(f"✓ Documento creado: {output_file.name}")
            print(f"  Tamaño: {size_mb:.2f} MB")
            print(f"  Imágenes totales: {total_images}")
            return True
        else:
            print(f"✗ Error: el archivo no se creó")
            return False

    except Exception as e:
        print(f"✗ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def verify_output(output_file):
    """Verifica el contenido del documento de salida."""
    try:
        doc = Document(str(output_file))

        print("\n" + "=" * 70)
        print("VERIFICACIÓN DEL DOCUMENTO FINAL")
        print("=" * 70)

        # Contar elementos
        total_paragraphs = len(doc.paragraphs)
        total_tables = len(doc.tables)
        total_images = count_images_in_docx(output_file)

        print(f"\n📄 Párrafos: {total_paragraphs:,}")
        print(f"📊 Tablas: {total_tables}")
        print(f"📸 Imágenes: {total_images}")

        # Tamaño
        file_size = output_file.stat().st_size
        print(f"💾 Tamaño: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)")

        print(f"\n✅ Verificación completa")

        return True
    except Exception as e:
        print(f"✗ Error en verificación: {e}")
        return False


def main():
    """Función principal."""
    print("=" * 70)
    print("Combinación de Documentos Word (usando docxcompose)")
    print("=" * 70)
    print(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print()

    # Verificar archivos
    print("Verificando archivos...")
    print("-" * 70)

    files_to_merge = []
    missing_files = []

    for filename in FILES_TO_MERGE:
        file_path = WORDS_DIR / filename
        if file_path.exists():
            size_kb = file_path.stat().st_size / 1024
            img_count = count_images_in_docx(file_path)
            img_info = f", {img_count} img" if img_count > 0 else ""
            print(f"  ✓ {filename} ({size_kb:.1f} KB{img_info})")
            files_to_merge.append(file_path)
        else:
            print(f"  ✗ {filename} - NO ENCONTRADO")
            missing_files.append(filename)

    print()

    if missing_files:
        print(f"✗ Error: Faltan {len(missing_files)} archivo(s)")
        return 1

    print(f"✓ Todos los archivos encontrados ({len(files_to_merge)} documentos)")
    print()

    # Combinar
    print("Combinando documentos...")
    print("-" * 70)

    if merge_documents(files_to_merge, OUTPUT_FILE):
        # Verificar
        verify_output(OUTPUT_FILE)

        print()
        print("=" * 70)
        print("✓ PROCESO COMPLETADO EXITOSAMENTE")
        print("=" * 70)
        print(f"\nDocumento final: {OUTPUT_FILE.name}")
        print(f"Ubicación: {OUTPUT_FILE}")
        print()
        return 0
    else:
        print()
        print("=" * 70)
        print("✗ ERROR EN EL PROCESO")
        print("=" * 70)
        print()
        return 1


if __name__ == "__main__":
    exit(main())
