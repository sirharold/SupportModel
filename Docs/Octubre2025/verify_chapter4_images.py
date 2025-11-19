#!/usr/bin/env python3
"""
Verifica si las imágenes del Capítulo 4 están en temp/capitulos.docx
"""

from docx import Document
from pathlib import Path

base_dir = Path("/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Octubre2025")
temp_dir = base_dir / "temp"
capitulos_path = temp_dir / "capitulos.docx"

print("VERIFICACIÓN DE IMÁGENES DEL CAPÍTULO 4")
print("=" * 70)

# Cargar documento de capítulos
doc = Document(str(capitulos_path))

# Contar imágenes
image_count = 0
image_list = []

for rel_id, rel in doc.part.rels.items():
    if "image" in rel.target_ref:
        image_count += 1
        image_list.append((rel_id, rel.target_ref))

print(f"\n🖼️  Total de imágenes en capitulos.docx: {image_count}\n")

if image_list:
    print("Lista de imágenes:")
    for i, (rel_id, target) in enumerate(image_list, 1):
        print(f"   {i}. {target} (ID: {rel_id})")

# Buscar referencias a Figuras 1-4 del Capítulo 4
print(f"\n📊 Buscando referencias a Figuras 1-4 en el texto...\n")

figura_refs = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if any(f"Figura {n}:" in text for n in [1, 2, 3, 4]):
        figura_refs.append((i, text[:100]))

print(f"Referencias a Figuras 1-4 encontradas: {len(figura_refs)}\n")
for i, (para_num, text) in enumerate(figura_refs, 1):
    print(f"   {i}. Párrafo {para_num}: {text}...")

# Buscar imágenes inline en los párrafos (usando runs)
print(f"\n🔍 Buscando imágenes inline en párrafos...\n")

inline_images = 0
for i, para in enumerate(doc.paragraphs):
    for run in para.runs:
        # Check if run has inline shapes (images)
        if hasattr(run._element, 'drawing_lst'):
            inline_images += 1
            print(f"   Imagen inline en párrafo {i}: {para.text[:50]}...")

print(f"\nTotal de imágenes inline encontradas: {inline_images}")

print(f"\n{'='*70}")
