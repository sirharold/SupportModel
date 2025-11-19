#!/usr/bin/env python3
"""
Combina documentos correctamente:
1. First.docx (portada) - numeración romana
2. preliminares.docx (Resumen + Abstract) - numeración romana continuada
3. capitulos.docx (Capítulos 1-8 + Bibliografía) - numeración arábiga desde 1

La numeración arábiga inicia en el Capítulo 1 (Introducción)
"""

from docx import Document
from docxcompose.composer import Composer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

base_dir = Path("/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Octubre2025")
words_dir = base_dir / "Words"
temp_dir = base_dir / "temp"

def set_page_number_format(section, format_type='decimal', start=1):
    """Establece el formato de numeración de páginas para una sección

    Args:
        section: Sección del documento
        format_type: 'decimal' (1,2,3) o 'lowerRoman' (i,ii,iii)
        start: Número de inicio
    """
    sectPr = section._sectPr

    # Remover elemento pgNumType existente si lo hay
    for child in list(sectPr):
        if child.tag == qn('w:pgNumType'):
            sectPr.remove(child)

    # Crear nuevo elemento pgNumType
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), format_type)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)

print("COMBINANDO DOCUMENTOS CON NUMERACIÓN CORRECTA")
print("=" * 70)

# 1. Cargar First.docx
first_path = words_dir / "First.docx"
print(f"\n1. Cargando: First.docx (portada)")
master = Document(first_path)
composer = Composer(master)

# Configurar numeración romana para First
print(f"   - Configurando numeración romana (i, ii, iii...)")
for section in master.sections:
    set_page_number_format(section, format_type='lowerRoman', start=1)

# 2. Agregar preliminares.docx (Resumen + Abstract)
prelim_path = temp_dir / "preliminares.docx"
print(f"\n2. Agregando: preliminares.docx (Resumen + Abstract)")
preliminares = Document(prelim_path)

# Configurar continuación de numeración romana en preliminares
for section in preliminares.sections:
    set_page_number_format(section, format_type='lowerRoman', start=1)

composer.append(preliminares)

# 3. Agregar salto de sección para cambiar a numeración arábiga
print(f"\n3. Insertando salto de sección para cambio de numeración")
master.add_section()

# Configurar numeración arábiga para la nueva sección (capítulos)
print(f"   - La próxima sección tendrá numeración arábiga desde 1")
last_section_idx = len(master.sections) - 1
set_page_number_format(master.sections[last_section_idx], format_type='decimal', start=1)

# 4. Agregar capitulos.docx (Capítulos 1-8 + Bibliografía)
cap_path = temp_dir / "capitulos.docx"
print(f"\n4. Agregando: capitulos.docx (Capítulos 1-8 + Bibliografía)")
capitulos = Document(cap_path)

# Remover numeración de página de capitulos para que herede de la sección anterior
for section in capitulos.sections:
    sectPr = section._sectPr
    for child in list(sectPr):
        if child.tag == qn('w:pgNumType'):
            sectPr.remove(child)

composer.append(capitulos)

# 5. Guardar documento combinado
output_path = words_dir / "ProyectoMagisterHaroldGomez.docx"
print(f"\n5. Guardando: {output_path.name}")
composer.save(str(output_path))

# Verificación
doc = Document(str(output_path))
image_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.target_ref)

print(f"\n{'='*70}")
print(f"✅ DOCUMENTO COMBINADO CREADO")
print(f"{'='*70}")
print(f"   Secciones: {len(doc.sections)}")
print(f"   Párrafos: {len(doc.paragraphs)}")
print(f"   Tablas: {len(doc.tables)}")
print(f"   Imágenes embebidas: {image_count}")
print(f"\n   Numeración de páginas:")
print(f"   - First + preliminares: Números romanos (i, ii, iii...)")
print(f"   - Capítulos 1-8 + Bibliografía: Números arábigos (1, 2, 3...)")
print(f"\n   Ubicación: {output_path}")
print(f"   Tamaño: {output_path.stat().st_size / (1024*1024):.1f} MB")
print(f"{'='*70}")
