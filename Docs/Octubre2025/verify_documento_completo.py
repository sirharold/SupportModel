#!/usr/bin/env python3
"""
Verifica imágenes en temp/documento_completo.docx
"""

from docx import Document
from pathlib import Path

base_dir = Path("/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Octubre2025")
temp_dir = base_dir / "temp"
doc_path = temp_dir / "documento_completo.docx"

print("VERIFICACIÓN DE temp/documento_completo.docx")
print("=" * 70)

doc = Document(str(doc_path))

# Contar imágenes
image_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.target_ref)

print(f"\n🖼️  Total de imágenes embebidas: {image_count}")
print(f"   Tamaño del archivo: {doc_path.stat().st_size / (1024*1024):.1f} MB")
print(f"   Párrafos: {len(doc.paragraphs)}")
print(f"   Tablas: {len(doc.tables)}")

print(f"\n{'='*70}")
