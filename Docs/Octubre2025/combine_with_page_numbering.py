#!/usr/bin/env python3
"""
Combina First.docx, preliminares.docx y capitulos.docx
Aplica numeración de páginas:
- First + preliminares: números romanos en minúscula (i, ii, iii, ...)
- Capítulos: números arábigos comenzando en 1 (1, 2, 3, ...)
"""

from docx import Document
from docxcompose.composer import Composer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

base_dir = Path("/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Octubre2025")
words_dir = base_dir / "Words"

def set_page_number_format(section, format_type='decimal', start=1):
    """Establece el formato de numeración de páginas para una sección

    Args:
        section: Sección del documento
        format_type: 'decimal' (1,2,3) o 'lowerRoman' (i,ii,iii)
        start: Número de inicio
    """
    sectPr = section._sectPr

    # Remover elemento pgNumType existente si lo hay
    for child in list(sectPr):
        if child.tag == qn('w:pgNumType'):
            sectPr.remove(child)

    # Crear nuevo elemento pgNumType
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), format_type)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)

print("COMBINANDO DOCUMENTOS CON NUMERACIÓN DE PÁGINAS")
print("=" * 70)

# 1. Cargar First.docx
first_path = words_dir / "First.docx"
print(f"\n1. Cargando: First.docx")
master = Document(first_path)
composer = Composer(master)

# 2. Agregar preliminares.docx (Resumen + Abstract)
prelim_path = base_dir / "preliminares.docx"
print(f"2. Agregando: preliminares.docx")
preliminares = Document(prelim_path)
composer.append(preliminares)

# 3. Configurar numeración romana para First + preliminares (secciones 0 y 1)
print(f"3. Configurando numeración romana (i, ii, iii) para preliminares")
for i in range(len(master.sections) - 1):  # Todas menos la última
    set_page_number_format(master.sections[i], format_type='lowerRoman', start=1)

# 4. Agregar salto de sección antes de los capítulos
print(f"4. Insertando salto de sección antes de capítulos")
last_para = master.paragraphs[-1]
last_para._element.addnext(last_para._element)  # Crear nuevo párrafo
master.add_section()

# 5. Agregar capitulos.docx
cap_path = base_dir / "capitulos.docx"
print(f"5. Agregando: capitulos.docx")
capitulos = Document(cap_path)
composer.append(capitulos)

# 6. Configurar numeración arábiga para capítulos (última sección)
print(f"6. Configurando numeración arábiga (1, 2, 3) para capítulos")
set_page_number_format(master.sections[-1], format_type='decimal', start=1)

# 7. Guardar documento combinado
output_path = words_dir / "ProyectoMAgisterHaroldGomez.docx"
print(f"\n7. Guardando: {output_path.name}")
composer.save(str(output_path))

# Estadísticas
doc = Document(str(output_path))
print(f"\n{'='*70}")
print(f"✅ DOCUMENTO COMBINADO CREADO")
print(f"{'='*70}")
print(f"   Secciones: {len(doc.sections)}")
print(f"   Párrafos: {len(doc.paragraphs)}")
print(f"   Tablas: {len(doc.tables)}")
print(f"\n   Numeración de páginas:")
print(f"   - Secciones 0-{len(master.sections)-2}: Números romanos (i, ii, iii)")
print(f"   - Sección {len(master.sections)-1}: Números arábigos (1, 2, 3)")
print(f"{'='*70}")
