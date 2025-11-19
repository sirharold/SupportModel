#!/usr/bin/env python3
"""
Combina First.docx con el contenido de la tesis usando docxcompose
"""

from docx import Document
from docxcompose.composer import Composer

print("Combinando documentos con docxcompose...")

# Cargar el documento base (First.docx)
master_doc = Document('First.docx')
composer = Composer(master_doc)

# Cargar el documento con los capítulos
chapters_doc = Document('ProyectoMAgisterHaroldGomez_temp.docx')

print(f"  First.docx: {len(master_doc.paragraphs)} párrafos")
print(f"  Capítulos: {len(chapters_doc.paragraphs)} párrafos")

# Verificar imágenes en el documento de capítulos
img_count = 0
for rel in chapters_doc.part.rels.values():
    if 'image' in rel.target_ref:
        img_count += 1
print(f"  Imágenes en capítulos: {img_count}")

# Combinar documentos
print("\nCombinando...")
composer.append(chapters_doc)

# Guardar el resultado
composer.save('ProyectoMAgisterHaroldGomez.docx')

# Verificar resultado
final_doc = Document('ProyectoMAgisterHaroldGomez.docx')
print(f"\n✓ Documento combinado guardado")
print(f"\n📄 Verificación:")
print(f"  Total de párrafos: {len(final_doc.paragraphs)}")
print(f"  Total de tablas: {len(final_doc.tables)}")

# Contar saltos de página
page_breaks = 0
for paragraph in final_doc.paragraphs:
    for run in paragraph.runs:
        if 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
            page_breaks += 1
print(f"  Saltos de página: {page_breaks}")

# Contar imágenes
image_count = 0
for rel in final_doc.part.rels.values():
    if 'image' in rel.target_ref:
        image_count += 1
print(f"  Total de imágenes: {image_count}")
