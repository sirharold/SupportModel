#!/usr/bin/env python3
"""
Combina First.docx con el contenido de la tesis preservando saltos de página
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import shutil

# Crear copia de First.docx como base
shutil.copy('First.docx', 'ProyectoMAgisterHaroldGomez.docx')

# Abrir el documento base y el documento con los capítulos
base_doc = Document('ProyectoMAgisterHaroldGomez.docx')
chapters_doc = Document('ProyectoMAgisterHaroldGomez_temp.docx')

print("Combinando documentos preservando saltos de página...")
print(f"  First.docx: {len(base_doc.paragraphs)} párrafos")
print(f"  Capítulos: {len(chapters_doc.paragraphs)} párrafos")

# Copiar todos los elementos del documento de capítulos
elements_copied = 0
page_breaks_copied = 0

for element in chapters_doc.element.body:
    # Clonar el elemento completo (preserva saltos de página y todo)
    base_doc.element.body.append(element)
    elements_copied += 1

    # Verificar si tiene salto de página
    if hasattr(element, 'xml') and 'w:br' in str(element.xml) and 'w:type="page"' in str(element.xml):
        page_breaks_copied += 1

print(f"\n✓ Elementos copiados: {elements_copied}")
print(f"✓ Saltos de página preservados: detectados en el proceso")

# Guardar el documento combinado
base_doc.save('ProyectoMAgisterHaroldGomez.docx')

# Verificar resultado
final_doc = Document('ProyectoMAgisterHaroldGomez.docx')
print(f"\n📄 Documento final:")
print(f"  Total de párrafos: {len(final_doc.paragraphs)}")
print(f"  Total de tablas: {len(final_doc.tables)}")

# Contar saltos de página en el resultado
page_breaks = 0
for paragraph in final_doc.paragraphs:
    for run in paragraph.runs:
        if 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
            page_breaks += 1

print(f"  Saltos de página verificados: {page_breaks}")

# Contar imágenes
image_count = 0
for rel in final_doc.part.rels.values():
    if 'image' in rel.target_ref:
        image_count += 1
print(f"  Total de imágenes: {image_count}")
