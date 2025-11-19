#!/usr/bin/env python3
"""
Combina los documentos con numeración de página correcta:
- First, Resumen, Abstract: números romanos (i, ii, iii)
- Introducción hasta Bibliografía: números arábigos (1, 2, 3...)
"""

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docxcompose.composer import Composer
import shutil

def set_page_number_format(section, format_type='decimal', start=1):
    """
    Configura el formato de numeración de página para una sección
    format_type: 'decimal' (1,2,3) o 'lowerRoman' (i,ii,iii)
    """
    sectPr = section._sectPr

    # Buscar o crear el elemento pgNumType
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)

    # Configurar formato y punto de inicio
    pgNumType.set(qn('w:fmt'), format_type)
    pgNumType.set(qn('w:start'), str(start))

def add_section_break(doc):
    """Agrega un salto de sección al final del documento"""
    last_paragraph = doc.paragraphs[-1]
    last_paragraph._element.get_or_add_pPr().get_or_add_sectPr()

print("Combinando documentos con numeración de páginas...")

# 1. Combinar First + preliminares (numeración romana)
print("\n1. Creando sección preliminar (numeración romana)...")
master_doc = Document('First.docx')
composer = Composer(master_doc)

# Agregar preliminares (Resumen + Abstract)
preliminares_doc = Document('preliminares.docx')
composer.append(preliminares_doc)

# Configurar todas las secciones hasta ahora con numeración romana
for i, section in enumerate(master_doc.sections):
    set_page_number_format(section, format_type='lowerRoman', start=1)
    print(f"  Sección {i}: numeración romana")

# Agregar un salto de sección al final de preliminares
# Esto creará una nueva sección para los capítulos
last_para = master_doc.paragraphs[-1]
new_section = master_doc.add_section(WD_SECTION.NEW_PAGE)
set_page_number_format(new_section, format_type='decimal', start=1)
print(f"  ✓ Salto de sección agregado para separar preliminares de capítulos")

# Guardar temporal
composer.save('temp_preliminares.docx')

# 2. Agregar capítulos (numeración arábiga)
print("\n2. Agregando capítulos (numeración arábiga)...")

# Cargar el documento preliminar
main_doc = Document('temp_preliminares.docx')
composer = Composer(main_doc)

# Contar secciones actuales (antes de agregar capítulos)
num_sections_before = len(main_doc.sections)
print(f"  Secciones en preliminares: {num_sections_before}")

# Agregar capítulos
capitulos_doc = Document('capitulos.docx')
composer.append(capitulos_doc)

# Configurar la primera sección de capítulos con numeración arábiga
# Esta es la sección que viene después de las preliminares
sections = main_doc.sections
print(f"  Total de secciones después de combinar: {len(sections)}")

# La última sección debería ser la de capítulos (ya configurada antes de combinar)
# Pero vamos a asegurarnos de que esté correctamente configurada
if num_sections_before < len(sections):
    # La sección de capítulos ya debería tener numeración arábiga
    print(f"  Verificando sección {num_sections_before-1} (capítulos)...")
    set_page_number_format(sections[num_sections_before-1], format_type='decimal', start=1)
    print(f"  ✓ Sección de capítulos configurada con numeración arábiga")

print("\n3. Guardando documento final...")
composer.save('ProyectoMAgisterHaroldGomez.docx')

# Verificación
final_doc = Document('ProyectoMAgisterHaroldGomez.docx')
print(f"\n✓ Documento creado:")
print(f"  Total de párrafos: {len(final_doc.paragraphs)}")
print(f"  Total de tablas: {len(final_doc.tables)}")
print(f"  Total de secciones: {len(final_doc.sections)}")

# Verificar imágenes
image_count = 0
for rel in final_doc.part.rels.values():
    if 'image' in rel.target_ref:
        image_count += 1
print(f"  Total de imágenes: {image_count}")

# Verificar saltos de página
page_breaks = 0
for paragraph in final_doc.paragraphs:
    for run in paragraph.runs:
        if 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
            page_breaks += 1
print(f"  Saltos de página: {page_breaks}")

# Mostrar configuración de numeración de cada sección
print(f"\n📄 Configuración de numeración por sección:")
for i, section in enumerate(final_doc.sections):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is not None:
        fmt = pgNumType.get(qn('w:fmt'), 'default')
        start = pgNumType.get(qn('w:start'), '1')
        print(f"  Sección {i}: formato={fmt}, inicio={start}")
    else:
        print(f"  Sección {i}: sin configuración de numeración")

print(f"\n✓ Esperado:")
print(f"  - First/Resumen/Abstract: números romanos (i, ii, iii)")
print(f"  - Capítulos 1-8/Bibliografía: números arábigos (1, 2, 3...)")

# Limpiar archivos temporales
import os
if os.path.exists('temp_preliminares.docx'):
    os.remove('temp_preliminares.docx')
    print(f"\n✓ Archivos temporales eliminados")
