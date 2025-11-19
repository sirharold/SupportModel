#!/usr/bin/env python3
"""
Verifica los saltos de página en el documento combinado
"""

from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.section import CT_SectPr

doc = Document('ProyectoMAgisterHaroldGomez.docx')

print("Analizando saltos de página y secciones...\n")

# Contar secciones
sections = len(doc.sections)
print(f"Total de secciones: {sections}")

# Buscar saltos de página
page_breaks = 0
section_breaks = 0

for i, paragraph in enumerate(doc.paragraphs):
    # Buscar saltos de página en el párrafo
    for run in paragraph.runs:
        if 'w:br' in run._element.xml:
            if 'w:type="page"' in run._element.xml:
                page_breaks += 1
                # Mostrar contexto del salto
                text = paragraph.text[:50] if paragraph.text else "[párrafo vacío]"
                print(f"  Salto de página #{page_breaks} en párrafo {i}: {text}...")

    # Verificar si el párrafo tiene salto de sección
    if paragraph._element.pPr is not None:
        sectPr = paragraph._element.pPr.find('.//{%s}sectPr' % paragraph._element.nsmap['w'])
        if sectPr is not None:
            section_breaks += 1
            text = paragraph.text[:50] if paragraph.text else "[párrafo vacío]"
            print(f"  Salto de sección #{section_breaks} en párrafo {i}: {text}...")

print(f"\n📊 Resumen:")
print(f"  - Saltos de página explícitos: {page_breaks}")
print(f"  - Saltos de sección: {section_breaks}")
print(f"  - Total de párrafos: {len(doc.paragraphs)}")
print(f"  - Total de secciones: {sections}")

# Verificar algunos encabezados principales
print(f"\n📄 Primeros encabezados de nivel 1:")
count = 0
for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith('Heading 1') or (para.text and para.text.startswith('#')):
        print(f"  Párrafo {i}: {para.text[:60]}")
        count += 1
        if count >= 15:
            break
