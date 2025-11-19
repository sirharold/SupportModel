#!/usr/bin/env python3
"""
Aplica formato de sangría y estilos a los párrafos del documento final
"""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("Aplicando formato de sangría a párrafos...")

doc = Document('ProyectoMAgisterHaroldGomez.docx')

# Configuración de sangría estándar para tesis (1.27 cm = 0.5 pulgadas)
SANGRIA_PRIMERA_LINEA = Cm(1.27)

# Contador
paragraphs_modified = 0
excluded_styles = ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
                   'Title', 'Subtitle', 'TOC', 'Caption', 'Header', 'Footer',
                   'List Paragraph', 'Quote', 'Intense Quote']

print(f"\nAplicando sangría de {SANGRIA_PRIMERA_LINEA.cm:.2f} cm a párrafos normales...\n")

for i, para in enumerate(doc.paragraphs):
    # Solo aplicar a párrafos con texto y que no sean títulos, citas, etc.
    if para.text.strip() and not any(style in para.style.name for style in excluded_styles):
        # Aplicar sangría de primera línea
        para.paragraph_format.first_line_indent = SANGRIA_PRIMERA_LINEA

        # Asegurar alineación justificada (estándar en tesis)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        paragraphs_modified += 1

        # Mostrar algunos ejemplos
        if paragraphs_modified <= 5:
            print(f"  ✓ Párrafo {i:4d}: {para.text[:60]}...")

print(f"\n✓ Total de párrafos modificados: {paragraphs_modified}")

# También configurar el estilo Normal para futuros párrafos
if 'Normal' in doc.styles:
    normal_style = doc.styles['Normal']
    normal_style.paragraph_format.first_line_indent = SANGRIA_PRIMERA_LINEA
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    print(f"✓ Estilo 'Normal' configurado con sangría de {SANGRIA_PRIMERA_LINEA.cm:.2f} cm")

# Guardar el documento
doc.save('ProyectoMAgisterHaroldGomez.docx')

print(f"\n✓ Documento guardado con formato aplicado")

# Verificación final
print(f"\n📄 Verificación:")
doc_verify = Document('ProyectoMAgisterHaroldGomez.docx')
with_indent = 0
for para in doc_verify.paragraphs:
    if para.paragraph_format.first_line_indent and para.paragraph_format.first_line_indent > 0:
        with_indent += 1

print(f"  Párrafos con sangría: {with_indent}")
print(f"  Total de párrafos: {len(doc_verify.paragraphs)}")
