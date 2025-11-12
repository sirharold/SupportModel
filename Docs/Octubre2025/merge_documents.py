#!/usr/bin/env python3
"""
Script para combinar múltiples archivos Word en un solo documento

Combina los siguientes archivos en orden:
1. First.docx (portada y contenido inicial)
2. Capitulo_0_Resumen.docx (resumen en español)
3. Capitulo_0_Abstract.docx (abstract en inglés)
4. Capitulo_1_Introduccion.docx
5. Capitulo_2_Estado_del_Arte.docx
6. Capitulo_3_Marco_Teorico.docx
7. Capitulo_4_Analisis_Exploratorio.docx
8. Capitulo_5_Metodologia.docx
9. Capitulo_6_Implementacion.docx

Resultado: ProyectoFinalHaroldGomez.docx

Uso:
    python merge_documents.py

Requisitos:
    pip install python-docx

Autor: Claude AI
Fecha: 2025-11-07
"""

import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configuración
WORDS_DIR = Path(__file__).parent / "Words"
OUTPUT_FILE = WORDS_DIR / "ProyectoFinalHaroldGomez.docx"

# Archivos a combinar en orden
FILES_TO_MERGE = [
    'First.docx',
    'Capitulo_0_Resumen.docx',
    'Capitulo_0_Abstract.docx',
    'Capitulo_1_Introduccion.docx',
    'Capitulo_2_Estado_del_Arte.docx',
    'Capitulo_3_Marco_Teorico.docx',
    'Capitulo_4_Analisis_Exploratorio.docx',
    'Capitulo_5_Metodologia.docx',
    'Capitulo_6_Implementacion.docx'
]


def add_section_break(doc):
    """Agrega un cambio de sección (section break) al documento."""
    # Obtener el último párrafo del documento
    last_paragraph = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()

    # Crear el elemento de propiedades de párrafo si no existe
    pPr = last_paragraph._element.get_or_add_pPr()

    # Crear el elemento de cambio de sección (nextPage)
    sectPr = OxmlElement('w:sectPr')

    # Tipo de cambio de sección: nextPage (nueva página)
    type_element = OxmlElement('w:type')
    type_element.set(qn('w:val'), 'nextPage')
    sectPr.append(type_element)

    # Agregar el cambio de sección al párrafo
    pPr.append(sectPr)


def remove_empty_paragraphs_at_end(doc, max_empty=10):
    """
    Remueve párrafos vacíos y saltos de página al final del documento.

    Args:
        doc: Documento de python-docx
        max_empty: Número máximo de elementos vacíos a remover desde el final
    """
    removed_count = 0
    removed_breaks = 0

    while doc.paragraphs and removed_count < max_empty:
        last_para = doc.paragraphs[-1]

        # Verificar si tiene saltos de página (page breaks)
        has_page_break = False
        for run in last_para.runs:
            if run._element.xml.find('w:br') != -1:
                # Hay un salto de página en este run
                has_page_break = True
                break

        # Verificar si el párrafo está vacío o solo tiene saltos
        if not last_para.text.strip() or has_page_break:
            # Remover el elemento
            p = last_para._element
            p.getparent().remove(p)
            if has_page_break:
                removed_breaks += 1
            removed_count += 1
        else:
            # Si encontramos un párrafo con contenido real, detenemos
            break

    if removed_count > 0:
        print(f"    → Removidos {removed_count} elementos vacíos al final ({removed_breaks} con saltos de página)")


def merge_documents(files, output_file):
    """
    Combina múltiples documentos Word en uno solo.

    Args:
        files: Lista de rutas a archivos .docx
        output_file: Ruta del archivo de salida

    Returns:
        True si fue exitoso, False en caso contrario
    """
    try:
        # Crear documento base con el primer archivo
        print(f"\n1. Cargando documento base: {files[0].name}")
        merged_doc = Document(files[0])

        # Agregar cada documento subsecuente
        for i, file_path in enumerate(files[1:], 2):
            print(f"{i}. Agregando: {file_path.name}")

            # Limpiar párrafos vacíos al final del documento actual
            remove_empty_paragraphs_at_end(merged_doc, max_empty=20)

            # Agregar cambio de sección antes del nuevo documento
            add_section_break(merged_doc)

            # Cargar el documento a agregar
            sub_doc = Document(file_path)

            # Copiar todos los elementos del documento
            for element in sub_doc.element.body:
                merged_doc.element.body.append(element)

        # Guardar el documento combinado
        print(f"\nGuardando documento combinado...")
        merged_doc.save(output_file)

        # Verificar que se creó
        if output_file.exists():
            size_mb = output_file.stat().st_size / (1024 * 1024)
            print(f"✓ Documento creado exitosamente: {output_file.name}")
            print(f"  Tamaño: {size_mb:.2f} MB")
            return True
        else:
            print(f"✗ Error: el archivo no se creó")
            return False

    except Exception as e:
        print(f"✗ Error al combinar documentos: {str(e)}")
        return False


def main():
    """Función principal del script."""
    print("=" * 70)
    print("Combinación de Documentos Word")
    print("=" * 70)
    print(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print()

    # Verificar que la carpeta Words existe
    if not WORDS_DIR.exists():
        print(f"✗ Error: La carpeta {WORDS_DIR} no existe")
        return 1

    print(f"Carpeta de trabajo: {WORDS_DIR}")
    print()

    # Verificar que todos los archivos existen
    print("Verificando archivos...")
    print("-" * 70)

    files_to_merge = []
    missing_files = []

    for filename in FILES_TO_MERGE:
        file_path = WORDS_DIR / filename
        if file_path.exists():
            size_kb = file_path.stat().st_size / 1024
            print(f"  ✓ {filename} ({size_kb:.1f} KB)")
            files_to_merge.append(file_path)
        else:
            print(f"  ✗ {filename} - NO ENCONTRADO")
            missing_files.append(filename)

    print()

    if missing_files:
        print(f"✗ Error: Faltan {len(missing_files)} archivo(s):")
        for f in missing_files:
            print(f"  - {f}")
        return 1

    print(f"✓ Todos los archivos encontrados ({len(files_to_merge)} documentos)")
    print()

    # Combinar documentos
    print("Combinando documentos...")
    print("-" * 70)

    if merge_documents(files_to_merge, OUTPUT_FILE):
        print()
        print("=" * 70)
        print("✓ PROCESO COMPLETADO EXITOSAMENTE")
        print("=" * 70)
        print(f"\nDocumento final: {OUTPUT_FILE.name}")
        print(f"Ubicación: {OUTPUT_FILE}")
        print()
        return 0
    else:
        print()
        print("=" * 70)
        print("✗ ERROR EN EL PROCESO")
        print("=" * 70)
        print()
        return 1


if __name__ == "__main__":
    exit(main())
