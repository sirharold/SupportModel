#!/usr/bin/env python3
"""
Combina documentos preservando TODAS las imágenes
Usa una estrategia diferente para evitar pérdida de imágenes
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import shutil

base_dir = Path("/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Octubre2025")
words_dir = base_dir / "Words"
temp_dir = base_dir / "temp"

def set_page_number_format(section, format_type='decimal', start=1):
    """Establece el formato de numeración de páginas para una sección"""
    sectPr = section._sectPr

    # Remover elemento pgNumType existente si lo hay
    for child in list(sectPr):
        if child.tag == qn('w:pgNumType'):
            sectPr.remove(child)

    # Crear nuevo elemento pgNumType
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), format_type)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)

def copy_paragraph(source_para, target_doc):
    """Copia un párrafo completo incluyendo formato e imágenes"""
    # Crear nuevo párrafo en el documento destino
    new_para = target_doc.add_paragraph()

    # Copiar el XML completo del párrafo (preserva todo el formato e imágenes)
    new_para._element.getparent().replace(new_para._element, source_para._element)

    return new_para

def copy_table(source_table, target_doc):
    """Copia una tabla completa"""
    # Agregar la tabla al documento destino
    # Usamos el elemento XML directamente para preservar todo el formato
    target_doc._element.body.append(source_table._element)

print("COMBINANDO DOCUMENTOS PRESERVANDO TODAS LAS IMÁGENES")
print("=" * 70)

# 1. Cargar First.docx como base
first_path = words_dir / "First.docx"
print(f"\n1. Cargando base: First.docx")
master = Document(str(first_path))

# 2. Cargar preliminares y agregar contenido
prelim_path = temp_dir / "preliminares.docx"
print(f"2. Copiando contenido de: preliminares.docx")
prelim_doc = Document(str(prelim_path))

for element in prelim_doc.element.body:
    master.element.body.append(element)

# 3. Configurar numeración romana para preliminares
print(f"3. Configurando numeración romana para preliminares")
for i in range(len(master.sections)):
    set_page_number_format(master.sections[i], format_type='lowerRoman', start=1)

# 4. Agregar salto de sección antes de capítulos
print(f"4. Insertando salto de sección para capítulos")
master.add_section()

# 5. Cargar capítulos y agregar contenido
cap_path = temp_dir / "capitulos.docx"
print(f"5. Copiando contenido de: capitulos.docx")
capitulos_doc = Document(str(cap_path))

# Copiar todas las relaciones (imágenes, etc.) del documento de capítulos
print(f"   - Copiando relaciones e imágenes...")
for rel_id, rel in capitulos_doc.part.rels.items():
    if rel_id not in master.part.rels:
        master.part.rels[rel_id] = rel

# Copiar todos los elementos del body (párrafos, tablas, etc.)
for element in capitulos_doc.element.body:
    master.element.body.append(element)

# 6. Configurar numeración arábiga para capítulos (última sección)
print(f"6. Configurando numeración arábiga para capítulos")
set_page_number_format(master.sections[-1], format_type='decimal', start=1)

# 7. Guardar documento combinado
output_path = words_dir / "ProyectoMagisterHaroldGomez.docx"
print(f"\n7. Guardando: {output_path.name}")
master.save(str(output_path))

# Verificar imágenes en el resultado
result_doc = Document(str(output_path))
image_count = sum(1 for rel in result_doc.part.rels.values() if "image" in rel.target_ref)

print(f"\n{'='*70}")
print(f"✅ DOCUMENTO COMBINADO CREADO")
print(f"{'='*70}")
print(f"   Secciones: {len(result_doc.sections)}")
print(f"   Párrafos: {len(result_doc.paragraphs)}")
print(f"   Tablas: {len(result_doc.tables)}")
print(f"   Imágenes: {image_count}")
print(f"\n   Numeración de páginas:")
print(f"   - Secciones 0-{len(master.sections)-2}: Números romanos (i, ii, iii)")
print(f"   - Sección {len(master.sections)-1}: Números arábigos (1, 2, 3)")
print(f"\n   Ubicación: {output_path}")
print(f"{'='*70}")
