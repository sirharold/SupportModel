#!/usr/bin/env python3
"""
Verifica el documento final ProyectoMagisterHaroldGomez.docx
"""

from docx import Document
from docx.shared import Cm
from pathlib import Path

base_dir = Path("/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Octubre2025")
words_dir = base_dir / "Words"
doc_path = words_dir / "ProyectoMagisterHaroldGomez.docx"

print("VERIFICACIÓN DEL DOCUMENTO FINAL")
print("=" * 70)

# Cargar documento
doc = Document(str(doc_path))

# 1. Información básica
print(f"\n📄 ARCHIVO: {doc_path.name}")
print(f"   Tamaño: {doc_path.stat().st_size / (1024*1024):.1f} MB")

# 2. Estructura
print(f"\n📊 ESTRUCTURA:")
print(f"   Secciones: {len(doc.sections)}")
print(f"   Párrafos: {len(doc.paragraphs)}")
print(f"   Tablas: {len(doc.tables)}")

# 3. Numeración de páginas
print(f"\n🔢 NUMERACIÓN DE PÁGINAS:")
for i, section in enumerate(doc.sections):
    sectPr = section._sectPr
    pgNumType = sectPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgNumType')

    if pgNumType is not None:
        fmt = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fmt')
        start = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}start')
        print(f"   Sección {i}: {fmt} (inicia en {start})")
    else:
        print(f"   Sección {i}: Sin formato definido")

# 4. Imágenes embebidas
print(f"\n🖼️  IMÁGENES:")
image_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.target_ref:
        image_count += 1

print(f"   Imágenes embebidas: {image_count}")

# 5. Sangría de párrafos
print(f"\n📏 SANGRÍA DE PÁRRAFOS:")

def get_all_table_paragraphs(doc):
    table_paragraphs = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    table_paragraphs.add(id(paragraph._element))
    return table_paragraphs

table_para_ids = get_all_table_paragraphs(doc)

stats = {
    'con_sangria': 0,
    'sin_sangria': 0,
    'en_tabla': 0
}

for para in doc.paragraphs:
    if not para.text.strip():
        continue

    para_id = id(para._element)

    if para_id in table_para_ids:
        stats['en_tabla'] += 1
    elif para.paragraph_format.first_line_indent and para.paragraph_format.first_line_indent > Cm(0):
        stats['con_sangria'] += 1
    else:
        stats['sin_sangria'] += 1

print(f"   Párrafos con sangría: {stats['con_sangria']}")
print(f"   Párrafos sin sangría (títulos, etc.): {stats['sin_sangria']}")
print(f"   Párrafos en tablas (sin sangría): {stats['en_tabla']}")

# 6. Tablas con captions
print(f"\n📋 TABLAS CON CAPTIONS:")
table_captions = 0
for para in doc.paragraphs:
    if 'Tabla' in para.text and ':' in para.text and para.text.strip().startswith('**Tabla'):
        table_captions += 1

print(f"   Captions de tablas encontrados: {table_captions}")

# 7. Figuras con captions
print(f"\n🎨 FIGURAS CON CAPTIONS:")
figure_captions = 0
for para in doc.paragraphs:
    if 'Figura' in para.text and ':' in para.text and 'Figura' in para.text[:20]:
        figure_captions += 1

print(f"   Captions de figuras encontrados: {figure_captions}")

# 8. Verificar saltos de página
print(f"\n📄 SALTOS DE PÁGINA:")
page_breaks = 0
for para in doc.paragraphs:
    if para._element.xpath('.//w:br[@w:type="page"]'):
        page_breaks += 1

print(f"   Saltos de página encontrados: {page_breaks}")

print(f"\n{'='*70}")
print(f"✅ VERIFICACIÓN COMPLETADA")
print(f"{'='*70}")
