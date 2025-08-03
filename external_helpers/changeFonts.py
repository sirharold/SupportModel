from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import os

def cambiar_fuente(doc_path, fuente_origen="Cambria", fuente_nueva="Arial", output_path=None):
    # Cargar documento
    doc = Document(doc_path)

    # Cambiar fuente en párrafos normales
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            font = run.font
            if font.name == fuente_origen or font.name is None:
                run.font.name = fuente_nueva
                rPr = run._element.rPr
                rFonts = rPr.rFonts
                rFonts.set(qn('w:eastAsia'), fuente_nueva)

    # Cambiar fuente dentro de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        if font.name == fuente_origen or font.name is None:
                            run.font.name = fuente_nueva
                            rPr = run._element.rPr
                            rFonts = rPr.rFonts
                            rFonts.set(qn('w:eastAsia'), fuente_nueva)

    # Guardar resultado
    if not output_path:
        root, ext = os.path.splitext(doc_path)
        output_path = f"{root}_Arial{ext}"
    doc.save(output_path)
    print(f"Documento guardado como: {output_path}")


myfile = "/Users/haroldgomez/Library/CloudStorage/OneDrive-UniversidadSanSebastian/Magister Data Science/DocFinal/ProyectoTituloMagisterHaroldGomez.docx"
# USO
cambiar_fuente(myfile, fuente_origen="Cambria", fuente_nueva="Arial")
