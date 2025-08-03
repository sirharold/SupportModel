#!/usr/bin/env python3
"""
Script para convertir el capítulo 7 actualizado a formato Word con formato específico
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re

def setup_document_styles(doc):
    """Configurar estilos del documento"""
    
    # Estilo para títulos de capítulo
    try:
        chapter_style = doc.styles['Heading 1']
    except KeyError:
        chapter_style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    
    chapter_style.font.name = 'Arial'
    chapter_style.font.size = Pt(16)
    chapter_style.font.bold = True
    chapter_style.paragraph_format.space_before = Pt(12)
    chapter_style.paragraph_format.space_after = Pt(6)
    
    # Estilo para subtítulos nivel 2
    try:
        heading2_style = doc.styles['Heading 2']
    except KeyError:
        heading2_style = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    
    heading2_style.font.name = 'Arial'
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True
    heading2_style.paragraph_format.space_before = Pt(10)
    heading2_style.paragraph_format.space_after = Pt(4)
    
    # Estilo para subtítulos nivel 3
    try:
        heading3_style = doc.styles['Heading 3']
    except KeyError:
        heading3_style = doc.styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    
    heading3_style.font.name = 'Arial'
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True
    heading3_style.paragraph_format.space_before = Pt(8)
    heading3_style.paragraph_format.space_after = Pt(3)
    
    # Estilo para subtítulos nivel 4
    try:
        heading4_style = doc.styles['Heading 4']
    except KeyError:
        heading4_style = doc.styles.add_style('Heading 4', WD_STYLE_TYPE.PARAGRAPH)
    
    heading4_style.font.name = 'Arial'
    heading4_style.font.size = Pt(11)
    heading4_style.font.bold = True
    heading4_style.paragraph_format.space_before = Pt(6)
    heading4_style.paragraph_format.space_after = Pt(2)
    
    # Estilo para texto normal
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.15
    
    # Estilo para código
    try:
        code_style = doc.styles['Code']
    except KeyError:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.line_spacing = 1.0
    # Agregar fondo gris claro
    from docx.oxml.shared import OxmlElement, qn
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), 'F5F5F5')  # Gris claro
    code_style._element.get_or_add_pPr().append(shading_elm)

def add_bullet_point(doc, text, indent_level=0):
    """Agregar punto de viñeta negro"""
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles['Normal']
    
    # Configurar indentación
    paragraph.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
    paragraph.paragraph_format.hanging_indent = Inches(0.25)
    
    # Agregar viñeta negra
    run = paragraph.add_run('• ')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    
    # Agregar texto
    run = paragraph.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    
    return paragraph

def process_markdown_line(doc, line, section_counters, in_code_block=False):
    """Procesar una línea de markdown y agregar al documento Word"""
    line_stripped = line.strip()
    
    # Manejar bloques de código
    if line_stripped.startswith('```python'):
        return 'start_code'
    elif line_stripped == '```' and in_code_block:
        return 'end_code'
    elif in_code_block:
        # Línea de código - preservar espacios originales
        paragraph = doc.add_paragraph(line.rstrip())
        paragraph.style = doc.styles['Code']
        return 'in_code'
    
    if not line_stripped:
        doc.add_paragraph()
        return 'normal'
    
    # Títulos con numeración automática
    if line_stripped.startswith('# '):
        # Capítulo principal
        title = line_stripped[2:].strip()
        if 'ACTUALIZACIÓN SECCIÓN 7.2' in title:
            title = "7. RESULTADOS Y ANÁLISIS (ACTUALIZADO)"
        paragraph = doc.add_paragraph(title)
        paragraph.style = doc.styles['Heading 1']
        section_counters['chapter'] = 7
        section_counters['section'] = 0
        section_counters['subsection'] = 0
        section_counters['subsubsection'] = 0
        return 'normal'
        
    elif line_stripped.startswith('## '):
        # Sección nivel 2
        title = line_stripped[3:].strip()
        section_counters['section'] += 1
        section_counters['subsection'] = 0
        section_counters['subsubsection'] = 0
        
        # Ajustar numeración específica
        if '7.2.1' in title:
            section_counters['section'] = 2
            section_counters['subsection'] = 1
        elif '7.3' in title:
            section_counters['section'] = 3
            section_counters['subsection'] = 0
        elif '7.6' in title:
            section_counters['section'] = 6
            section_counters['subsection'] = 0
        
        # Remover numeración existente si existe
        title = re.sub(r'^7\.\d+(\.\d+)?\s*', '', title)
        
        numbered_title = f"7.{section_counters['section']} {title}"
        paragraph = doc.add_paragraph(numbered_title)
        paragraph.style = doc.styles['Heading 2']
        return 'normal'
        
    elif line_stripped.startswith('### '):
        # Subsección nivel 3
        title = line[4:].strip()
        section_counters['subsection'] += 1
        section_counters['subsubsection'] = 0
        
        # Remover numeración existente si existe
        title = re.sub(r'^7\.\d+\.\d+(\.\d+)?\s*', '', title)
        
        numbered_title = f"7.{section_counters['section']}.{section_counters['subsection']} {title}"
        paragraph = doc.add_paragraph(numbered_title)
        paragraph.style = doc.styles['Heading 3']
        
    elif line.startswith('#### '):
        # Subsubsección nivel 4
        title = line[5:].strip()
        section_counters['subsubsection'] += 1
        
        # Remover numeración existente si existe
        title = re.sub(r'^7\.\d+\.\d+\.\d+(\.\d+)?\s*', '', title)
        
        numbered_title = f"7.{section_counters['section']}.{section_counters['subsection']}.{section_counters['subsubsection']} {title}"
        paragraph = doc.add_paragraph(numbered_title)
        paragraph.style = doc.styles['Heading 4']
        
    elif line.startswith('- '):
        # Punto de viñeta
        text = line[2:].strip()
        # Remover formato de markdown para negritas
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        add_bullet_point(doc, text)
        
    else:
        # Texto normal
        if line:
            # Procesar texto con negritas
            paragraph = doc.add_paragraph()
            paragraph.style = doc.styles['Normal']
            
            # Dividir por negritas
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for i, part in enumerate(parts):
                if i % 2 == 0:  # Texto normal
                    if part:
                        run = paragraph.add_run(part)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                else:  # Texto en negrita
                    run = paragraph.add_run(part)
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.bold = True

def convert_chapter7_to_word():
    """Convertir el capítulo 7 actualizado a Word"""
    
    # Leer archivo markdown
    input_file = '/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Finales/capitulo_7_completo_actualizado.md'
    output_file = '/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Finales/words/capitulo_7_completo_actualizado.docx'
    
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        print(f"❌ Error: No se encontró el archivo {input_file}")
        return False
    
    # Crear documento Word
    doc = Document()
    
    # Configurar estilos
    setup_document_styles(doc)
    
    # Contadores para numeración automática
    section_counters = {
        'chapter': 7,
        'section': 0,
        'subsection': 0,
        'subsubsection': 0
    }
    
    # Procesar línea por línea
    lines = content.split('\n')
    for line in lines:
        process_markdown_line(doc, line, section_counters)
    
    # Guardar documento
    try:
        doc.save(output_file)
        print(f"✅ Capítulo 7 convertido exitosamente a: {output_file}")
        return True
    except Exception as e:
        print(f"❌ Error guardando el archivo: {e}")
        return False

if __name__ == "__main__":
    success = convert_chapter7_to_word()
    if success:
        print("🎉 Conversión completada!")
        print("📋 Características del documento:")
        print("  • Fuente: Arial en todo el documento")
        print("  • Numeración: Capítulo 7 con subsecciones automáticas")
        print("  • Viñetas: Puntos negros (•) en lugar de guiones")
        print("  • Formato: Espaciado y estilos profesionales")
    else:
        print("❌ Error en la conversión")