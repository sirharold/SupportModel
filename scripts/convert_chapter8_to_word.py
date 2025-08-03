"""
Script para convertir el Capítulo 8 actualizado a formato Word
Con numeración comenzando en 8, fuente Arial y bullets negros
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re
from pathlib import Path

def setup_styles(doc):
    """Configurar estilos del documento"""
    
    # Estilo para el título del capítulo
    try:
        chapter_title_style = doc.styles['Heading 1']
    except KeyError:
        chapter_title_style = doc.styles.add_style('Chapter Title', WD_STYLE_TYPE.PARAGRAPH)
    
    chapter_title_style.font.name = 'Arial'
    chapter_title_style.font.size = Pt(18)
    chapter_title_style.font.bold = True
    chapter_title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    chapter_title_style.paragraph_format.space_before = Pt(0)
    chapter_title_style.paragraph_format.space_after = Pt(18)
    
    # Estilos para subtítulos (niveles 2-6)
    heading_sizes = [16, 14, 12, 11, 10]
    
    for i, size in enumerate(heading_sizes, 2):
        try:
            heading_style = doc.styles[f'Heading {i}']
        except KeyError:
            heading_style = doc.styles.add_style(f'Heading {i}', WD_STYLE_TYPE.PARAGRAPH)
        
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(size)
        heading_style.font.bold = True
        heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(8)
    
    # Estilo para párrafos normales
    try:
        normal_style = doc.styles['Normal']
    except KeyError:
        normal_style = doc.styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
    
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.space_before = Pt(0)
    
    # Estilo para listas
    try:
        list_style = doc.styles['List Paragraph']
    except KeyError:
        list_style = doc.styles.add_style('List Paragraph', WD_STYLE_TYPE.PARAGRAPH)
    
    list_style.font.name = 'Arial'
    list_style.font.size = Pt(11)
    list_style.paragraph_format.left_indent = Inches(0.25)
    list_style.paragraph_format.space_after = Pt(3)
    
    # Estilo para código
    try:
        code_style = doc.styles['Code']
    except KeyError:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

def create_bullet_paragraph(doc, text, level=0):
    """Crear párrafo con bullet negro"""
    paragraph = doc.add_paragraph()
    paragraph.style = 'List Paragraph'
    
    # Configurar bullet negro
    paragraph.paragraph_format.left_indent = Inches(0.25 + (level * 0.25))
    
    # Crear elemento de lista con bullet negro
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    
    # Bullet level
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)
    
    # Bullet ID (usar ID estándar para bullets negros)
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')
    numPr.append(numId)
    
    pPr.append(numPr)
    
    # Agregar texto
    run = paragraph.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    
    return paragraph

def create_bullet_paragraph_with_formatting(doc, text, level=0):
    """Crear párrafo con bullet negro y formato markdown"""
    paragraph = doc.add_paragraph()
    paragraph.style = 'List Paragraph'
    
    # Configurar bullet negro
    paragraph.paragraph_format.left_indent = Inches(0.25 + (level * 0.25))
    
    # Crear elemento de lista con bullet negro
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    
    # Bullet level
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)
    
    # Bullet ID (usar ID estándar para bullets negros)
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')
    numPr.append(numId)
    
    pPr.append(numPr)
    
    # Procesar texto con formato
    add_formatted_text_to_paragraph(paragraph, text)
    
    return paragraph

def add_formatted_text_to_paragraph(paragraph, text):
    """Agregar texto con formato markdown a un párrafo"""
    # Usar regex más específico y procesamiento secuencial
    import re
    
    # Primero procesar negrita (**texto**)
    parts = []
    remaining_text = text
    
    # Encontrar todos los textos en negrita
    bold_pattern = r'\*\*(.*?)\*\*'
    bold_matches = list(re.finditer(bold_pattern, remaining_text))
    
    if bold_matches:
        last_end = 0
        for match in bold_matches:
            # Texto antes del bold
            if match.start() > last_end:
                parts.append(('normal', remaining_text[last_end:match.start()]))
            # Texto en bold
            parts.append(('bold', match.group(1)))
            last_end = match.end()
        
        # Texto después del último bold
        if last_end < len(remaining_text):
            parts.append(('normal', remaining_text[last_end:]))
    else:
        parts.append(('normal', remaining_text))
    
    # Ahora procesar cada parte para cursiva
    final_parts = []
    for part_type, part_text in parts:
        if part_type == 'bold':
            final_parts.append(('bold', part_text))
        else:
            # Procesar cursiva en texto normal
            italic_pattern = r'(?<!\*)\*([^*]+)\*(?!\*)'
            italic_matches = list(re.finditer(italic_pattern, part_text))
            
            if italic_matches:
                last_end = 0
                for match in italic_matches:
                    # Texto antes del italic
                    if match.start() > last_end:
                        final_parts.append(('normal', part_text[last_end:match.start()]))
                    # Texto en italic
                    final_parts.append(('italic', match.group(1)))
                    last_end = match.end()
                
                # Texto después del último italic
                if last_end < len(part_text):
                    final_parts.append(('normal', part_text[last_end:]))
            else:
                final_parts.append(('normal', part_text))
    
    # Procesar código inline en todas las partes normales
    final_final_parts = []
    for part_type, part_text in final_parts:
        if part_type in ['bold', 'italic']:
            final_final_parts.append((part_type, part_text))
        else:
            # Procesar código inline
            code_pattern = r'`([^`]+)`'
            code_matches = list(re.finditer(code_pattern, part_text))
            
            if code_matches:
                last_end = 0
                for match in code_matches:
                    # Texto antes del código
                    if match.start() > last_end:
                        final_final_parts.append(('normal', part_text[last_end:match.start()]))
                    # Texto código
                    final_final_parts.append(('code', match.group(1)))
                    last_end = match.end()
                
                # Texto después del último código
                if last_end < len(part_text):
                    final_final_parts.append(('normal', part_text[last_end:]))
            else:
                final_final_parts.append(('normal', part_text))
    
    # Agregar todas las partes al párrafo
    for part_type, part_text in final_final_parts:
        if not part_text:
            continue
            
        run = paragraph.add_run(part_text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        
        if part_type == 'bold':
            run.bold = True
        elif part_type == 'italic':
            run.italic = True
        elif part_type == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

def add_bullet_numbering(doc):
    """Agregar definición de numeración para bullets negros"""
    
    # Crear elemento de numeración
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        numbering_part = doc.part.add_numbering_part()
    
    # XML para bullet negro
    bullet_xml = f'''
    <w:abstractNum {nsdecls('w')} w:abstractNumId="0">
        <w:multiLevelType w:val="singleLevel"/>
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/>
            <w:numFmt w:val="bullet"/>
            <w:lvlText w:val="•"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="360" w:hanging="360"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:start w:val="1"/>
            <w:numFmt w:val="bullet"/>
            <w:lvlText w:val="•"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="720" w:hanging="360"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>
            </w:rPr>
        </w:lvl>
    </w:abstractNum>
    '''
    
    try:
        abstract_num = parse_xml(bullet_xml)
        numbering_part.element.append(abstract_num)
        
        # Crear instancia de numeración
        num_xml = f'''
        <w:num {nsdecls('w')} w:numId="1">
            <w:abstractNumId w:val="0"/>
        </w:num>
        '''
        num = parse_xml(num_xml)
        numbering_part.element.append(num)
    except Exception as e:
        print(f"Warning: No se pudo configurar bullets personalizados: {e}")

def process_markdown_line(doc, line, in_code_block=False):
    """Procesar una línea de markdown y convertirla a Word"""
    line = line.rstrip()
    
    # Detectar bloques de código
    if line.startswith('```'):
        return not in_code_block, None
    
    # Si estamos en un bloque de código
    if in_code_block:
        p = doc.add_paragraph(line)
        p.style = 'Code'
        # Agregar fondo gris
        try:
            shading = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w')))
            p._element.get_or_add_pPr().append(shading)
        except:
            pass
        return in_code_block, None
    
    # Líneas vacías - NO agregar párrafos vacíos
    if not line.strip():
        return in_code_block, 'skip'
    
    # Títulos
    if line.startswith('#'):
        level = len(line) - len(line.lstrip('#'))
        title_text = line.lstrip('#').strip()
        
        # Ajustar numeración del capítulo
        if level == 1 and title_text.startswith('8.'):
            heading_level = 1
        elif level == 2:
            heading_level = 2
        elif level == 3:
            heading_level = 3
        elif level == 4:
            heading_level = 4
        elif level == 5:
            heading_level = 5
        else:
            heading_level = 6
            
        p = doc.add_paragraph(title_text)
        p.style = f'Heading {heading_level}'
        return in_code_block, None
    
    # Listas con bullets (solo guiones, no asteriscos de formato)
    if line.strip().startswith('- '):
        # Determinar nivel de indentación
        indent_level = (len(line) - len(line.lstrip())) // 2
        bullet_text = line.strip()[2:].strip()  # Remover '- '
        
        # Crear párrafo con formato
        create_bullet_paragraph_with_formatting(doc, bullet_text, indent_level)
        return in_code_block, None
    
    # Párrafos normales
    p = doc.add_paragraph()
    p.style = 'Normal'
    
    # Procesar texto con formato
    text = line.strip()
    add_formatted_text_to_paragraph(p, text)
    
    return in_code_block, None

def convert_chapter8_to_word():
    """Convertir capítulo 8 actualizado a Word"""
    
    # Rutas de archivos
    input_file = "/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Finales/capitulo_8_actualizado.md"
    output_dir = Path("/Users/haroldgomez/Documents/ProyectoTituloMAgister/SupportModel/Docs/Finales/words")
    output_file = output_dir / "capitulo_8_actualizado.docx"
    
    # Crear directorio si no existe
    output_dir.mkdir(exist_ok=True)
    
    print(f"Convirtiendo: {input_file}")
    print(f"Destino: {output_file}")
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos
    setup_styles(doc)
    
    # Agregar numeración para bullets
    add_bullet_numbering(doc)
    
    # Leer archivo markdown
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print(f"Error: No se encontró el archivo {input_file}")
        return False
    except Exception as e:
        print(f"Error leyendo archivo: {e}")
        return False
    
    print(f"Procesando {len(lines)} líneas...")
    
    # Procesar líneas
    in_code_block = False
    
    for i, line in enumerate(lines):
        try:
            in_code_block, action = process_markdown_line(doc, line, in_code_block)
            # Skip líneas vacías para evitar espacios extra
            if action == 'skip':
                continue
        except Exception as e:
            print(f"Error procesando línea {i+1}: {e}")
            print(f"Línea: {line[:100]}...")
            continue
    
    # Guardar documento
    try:
        doc.save(str(output_file))
        print(f"✅ Conversión exitosa: {output_file}")
        return True
    except Exception as e:
        print(f"Error guardando documento: {e}")
        return False

def main():
    """Función principal"""
    print("=== CONVERSIÓN CAPÍTULO 8 A WORD ===")
    print("Configuración:")
    print("- Fuente: Arial")
    print("- Numeración: Comienza en 8")
    print("- Bullets: Negros (•)")
    print("- Formato: Justificado")
    print()
    
    success = convert_chapter8_to_word()
    
    if success:
        print("\n✅ CONVERSIÓN COMPLETADA EXITOSAMENTE")
        print("\nArchivo creado:")
        print("- capitulo_8_actualizado.docx")
        print("\nCaracterísticas del documento:")
        print("- Fuente Arial en todo el documento")
        print("- Títulos numerados comenzando en 8")
        print("- Bullets negros para listas")
        print("- Código con fuente Consolas y fondo gris")
        print("- Párrafos justificados")
    else:
        print("\n❌ ERROR EN LA CONVERSIÓN")
        print("Revisar mensajes de error anteriores")

if __name__ == "__main__":
    main()